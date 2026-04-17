"""
Conversor de Arquivos AM v2.0.0 - ExpansÃ£o Convertio-like
Ferramenta interna de conversÃ£o de arquivos para escritÃ³rio Andrade Maia.

MELHORIAS v2.0.0:
- ðŸŽ¨ Interface renovada com tema preto/laranja do escritÃ³rio
- ðŸ“Š Barra de progresso em tempo real
- â¬‡ï¸ Download integrado via AJAX
- ðŸ“± Design responsivo
- âš¡ Feedback visual aprimorado

IMPORTANTE: Esta aplicaÃ§Ã£o NÃƒO inclui funcionalidade de PDF.
A ferramenta PDF (I'AM.pdf) Ã© mantida completamente separada.

Arquitetura modular inspirada no I'AM.pdf, reutilizando padrÃµes de:
- Helpers genÃ©ricos (file_sha1, cache_path, save_converted_copy)
- Tratamento de uploads/downloads
- ValidaÃ§Ãµes de seguranÃ§a
- Estrutura Flask com rotas modulares

MÃ³dulos:
- conversores/: Classes especÃ­ficas para cada tipo de conversÃ£o (Ã¡udio, vÃ­deo, imagem)
- rotas/: Rotas Flask/FastAPI
- utils/: FunÃ§Ãµes utilitÃ¡rias reutilizÃ¡veis
- templates/: HTML frontend
- static/: CSS/JS
"""

import os
import json
import tempfile
import hashlib
import re
import logging
import shutil
from abc import ABC, abstractmethod
from typing import Optional, Tuple
from datetime import datetime
import sys
from pathlib import Path

from flask import Flask
try:
    from flask_sock import Sock
except ImportError:
    Sock = None
from config import get_config

# ConfiguraÃ§Ãµes
config = get_config()

app = Flask(__name__)
app.config.from_object(config)
sock = Sock(app) if Sock else None

# ========== CONFIGURAÃ‡ÃƒO DE FFmpeg ==========

def setup_ffmpeg_path():
    """Detecta e configura o caminho do FFmpeg."""
    ffmpeg_paths = [
        # Procurar localmente em ./ffmpeg/bin/
        os.path.join(os.path.dirname(__file__), "ffmpeg", "bin"),
        # Procurar em ./ffmpeg-master-latest-win64-gpl/bin/
        os.path.join(os.path.dirname(__file__), "ffmpeg-master-latest-win64-gpl", "bin"),
    ]
    
    # Adicionar ao PATH se existir
    for path in ffmpeg_paths:
        if os.path.exists(path):
            if path not in os.environ.get('PATH', ''):
                os.environ['PATH'] = path + os.pathsep + os.environ.get('PATH', '')
            return True
    
    return False

# Configurar FFmpeg ao iniciar
setup_ffmpeg_path()

# ========== CLASSE BASE PARA CONVERSORES ==========

class BaseConverter(ABC):
    """
    Classe base abstrata para conversores de arquivo.
    Segue padrÃ£o similar ao I'AM.pdf, com mÃ©todos utilitÃ¡rios compartilhados.
    """

    def __init__(self, input_path: str, output_format: str):
        self.input_path = input_path
        self.output_format = output_format
        self.temp_files = []  # Para limpeza

    @abstractmethod
    def convert(self) -> bytes:
        """
        MÃ©todo abstrato para conversÃ£o.
        Deve retornar os bytes do arquivo convertido.
        """
        pass

    def _get_output_filename(self, original_name: str) -> str:
        """Gera nome do arquivo de saÃ­da baseado no original."""
        stem, _ = os.path.splitext(original_name)
        return f"{stem}_convertido.{self.output_format}"

    def _create_temp_file(self, suffix: str = "") -> str:
        """Cria arquivo temporÃ¡rio e registra para limpeza."""
        fd, path = tempfile.mkstemp(suffix=suffix)
        os.close(fd)
        self.temp_files.append(path)
        return path

    def cleanup(self):
        """Remove arquivos temporÃ¡rios."""
        for path in self.temp_files:
            try:
                os.remove(path)
            except:
                pass
        self.temp_files = []

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        self.cleanup()

# ========== CONVERSORES ESPECÃFICOS ==========

class AudioConverter(BaseConverter):
    """Conversor de Ã¡udio: OGG â†’ MP3"""

    def convert(self) -> bytes:
        try:
            import ffmpeg
        except ImportError:
            raise RuntimeError("ffmpeg-python nÃ£o instalado. Instale com: pip install ffmpeg-python")

        output_path = self._create_temp_file(".mp3")

        # ConversÃ£o usando ffmpeg
        stream = ffmpeg.input(self.input_path)
        stream = ffmpeg.output(stream, output_path, acodec='mp3', audio_bitrate='192k')
        ffmpeg.run(stream, overwrite_output=True, quiet=True)

        with open(output_path, 'rb') as f:
            return f.read()

class VideoConverter(BaseConverter):
    """Conversor de vÃ­deo: MP4 â†’ WAV"""

    def convert(self) -> bytes:
        try:
            import ffmpeg
        except ImportError:
            raise RuntimeError("ffmpeg-python nÃ£o instalado. Instale com: pip install ffmpeg-python")

        output_path = self._create_temp_file(".wav")

        # Extrair Ã¡udio do vÃ­deo para WAV
        stream = ffmpeg.input(self.input_path)
        stream = ffmpeg.output(stream, output_path, acodec='pcm_s16le', ar='44100')
        ffmpeg.run(stream, overwrite_output=True, quiet=True)

        with open(output_path, 'rb') as f:
            return f.read()

class ImageConverter(BaseConverter):
    """Conversor de imagem: PNG â†’ JPG"""

    def convert(self) -> bytes:
        try:
            from PIL import Image
        except ImportError:
            raise RuntimeError("Pillow nÃ£o instalado. Instale com: pip install Pillow")

        # Abrir imagem e converter
        with Image.open(self.input_path) as img:
            # Converter para RGB se necessÃ¡rio (para JPG)
            if img.mode in ("RGBA", "LA", "P"):
                img = img.convert("RGB")

            output_path = self._create_temp_file(".jpg")
            img.save(output_path, 'JPEG', quality=90)

            with open(output_path, 'rb') as f:
                return f.read()

class DocumentConverter(BaseConverter):
    """Conversor de documentos: DOCX â†’ PDF"""

    def convert(self) -> bytes:
        try:
            from docx import Document
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        except ImportError:
            raise RuntimeError("python-docx e reportlab nÃ£o instalados")

        if self.output_format == 'pdf':
            # DOCX â†’ PDF
            doc = Document(self.input_path)
            output_path = self._create_temp_file(".pdf")

            # Criar PDF
            pdf = SimpleDocTemplate(output_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []

            for para in doc.paragraphs:
                if para.text.strip():
                    story.append(Paragraph(para.text, styles['Normal']))
                    story.append(Spacer(1, 12))

            pdf.build(story)

            with open(output_path, 'rb') as f:
                return f.read()

        elif self.output_format == 'docx':
            # PDF â†’ DOCX
            try:
                from pdf2docx import Converter
            except ImportError:
                raise RuntimeError("pdf2docx nÃ£o instalado")

            output_path = self._create_temp_file(".docx")

            cv = Converter(self.input_path)
            cv.convert(output_path, start=0, end=None)
            cv.close()

            with open(output_path, 'rb') as f:
                return f.read()

        elif self.output_format == 'txt':
            # DOCX â†’ TXT
            doc = Document(self.input_path)
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

            return text.encode('utf-8')

class TextConverter(BaseConverter):
    """Conversor de texto: TXT â†’ PDF"""

    def convert(self) -> bytes:
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        except ImportError:
            raise RuntimeError("reportlab nÃ£o instalado")

        # Tentar ler com diferentes codificaÃ§Ãµes
        encodings_to_try = ['utf-8', 'utf-16', 'utf-16-le', 'utf-16-be', 'latin-1', 'cp1252']
        text = None

        for encoding in encodings_to_try:
            try:
                with open(self.input_path, 'r', encoding=encoding) as f:
                    text = f.read()
                break
            except UnicodeDecodeError:
                continue

        if text is None:
            # Se nenhuma codificaÃ§Ã£o funcionou, tentar como bytes e decodificar
            with open(self.input_path, 'rb') as f:
                raw_bytes = f.read()
                # Remover BOM se existir
                if raw_bytes.startswith(b'\xff\xfe'):
                    raw_bytes = raw_bytes[2:]  # UTF-16 LE BOM
                elif raw_bytes.startswith(b'\xfe\xff'):
                    raw_bytes = raw_bytes[2:]  # UTF-16 BE BOM

                try:
                    text = raw_bytes.decode('utf-16-le')
                except:
                    try:
                        text = raw_bytes.decode('latin-1')
                    except:
                        raise ValueError("NÃ£o foi possÃ­vel decodificar o arquivo de texto")

        output_path = self._create_temp_file(".pdf")

        # Criar PDF
        pdf = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        for line in text.split('\n'):
            if line.strip():
                story.append(Paragraph(line, styles['Normal']))
                story.append(Spacer(1, 6))

        pdf.build(story)

        with open(output_path, 'rb') as f:
            return f.read()

class SpreadsheetConverter(BaseConverter):
    """Conversor de planilhas: XLSX â†’ CSV"""

    def convert(self) -> bytes:
        try:
            import pandas as pd
        except ImportError:
            raise RuntimeError("pandas nÃ£o instalado")

        if self.output_format == 'csv':
            # XLSX â†’ CSV
            df = pd.read_excel(self.input_path)
            csv_content = df.to_csv(index=False)
            return csv_content.encode('utf-8')

        elif self.output_format == 'xlsx':
            # CSV â†’ XLSX
            df = pd.read_csv(self.input_path)
            output_path = self._create_temp_file(".xlsx")
            df.to_excel(output_path, index=False)

            with open(output_path, 'rb') as f:
                return f.read()

class ImageWebPConverter(BaseConverter):
    """Conversor de imagem para WebP"""

    def convert(self) -> bytes:
        try:
            from PIL import Image
        except ImportError:
            raise RuntimeError("Pillow nÃ£o instalado")

        with Image.open(self.input_path) as img:
            # Converter para RGB se necessÃ¡rio
            if img.mode in ("RGBA", "LA", "P"):
                img = img.convert("RGB")

            output_path = self._create_temp_file(".webp")
            img.save(output_path, 'WEBP', quality=85)

            with open(output_path, 'rb') as f:
                return f.read()

class AudioWAVConverter(BaseConverter):
    """Conversor de Ã¡udio: MP3 â†’ WAV"""

    def convert(self) -> bytes:
        try:
            import ffmpeg
        except ImportError:
            raise RuntimeError("ffmpeg-python nÃ£o instalado")

        output_path = self._create_temp_file(".wav")

        # MP3 â†’ WAV
        stream = ffmpeg.input(self.input_path)
        stream = ffmpeg.output(stream, output_path, acodec='pcm_s16le', ar='44100')
        ffmpeg.run(stream, overwrite_output=True, quiet=True)

        with open(output_path, 'rb') as f:
            return f.read()

class VideoAVIConverter(BaseConverter):
    """Conversor de vÃ­deo: MP4 â†’ AVI"""

    def convert(self) -> bytes:
        try:
            import ffmpeg
        except ImportError:
            raise RuntimeError("ffmpeg-python nÃ£o instalado")

        output_path = self._create_temp_file(".avi")

        # MP4 â†’ AVI
        stream = ffmpeg.input(self.input_path)
        stream = ffmpeg.output(stream, output_path, vcodec='libx264', acodec='mp3')
        ffmpeg.run(stream, overwrite_output=True, quiet=True)

        with open(output_path, 'rb') as f:
            return f.read()

class VideoM4VConverter(BaseConverter):
    """Conversor de vÃ­deo: M4V â†’ MP4"""

    def convert(self) -> bytes:
        try:
            import ffmpeg
        except ImportError:
            raise RuntimeError("ffmpeg-python nÃ£o instalado")

        output_path = self._create_temp_file(".mp4")

        # M4V â†’ MP4 (re-encode com H.264/AAC)
        stream = ffmpeg.input(self.input_path)
        stream = ffmpeg.output(stream, output_path, vcodec='libx264', acodec='aac', preset='medium')
        ffmpeg.run(stream, overwrite_output=True, quiet=True)

        with open(output_path, 'rb') as f:
            return f.read()

# ========== CONVERSORES GENÃ‰RICOS (FFmpeg) ==========

class GenericAudioConverter(BaseConverter):
    """Conversor genÃ©rico de Ã¡udio para qualquer formato FFmpeg"""
    
    # Mapeamento de formatos para codec e parÃ¢metros FFmpeg
    AUDIO_FORMATS = {
        'mp3': {'acodec': 'mp3', 'audio_bitrate': '192k'},
        'wav': {'acodec': 'pcm_s16le', 'ar': '44100'},
        'flac': {'acodec': 'flac'},
        'aac': {'acodec': 'aac', 'audio_bitrate': '192k'},
        'm4a': {'acodec': 'aac', 'audio_bitrate': '192k'},
        'ogg': {'acodec': 'libvorbis', 'audio_bitrate': '192k'},
        'opus': {'acodec': 'libopus', 'audio_bitrate': '128k'},
        'aiff': {'acodec': 'pcm_s16le', 'f': 'aiff'},
        'weba': {'acodec': 'libopus', 'audio_bitrate': '128k'},
        'ac3': {'acodec': 'ac3', 'audio_bitrate': '192k'},
        'dts': {'acodec': 'dts', 'audio_bitrate': '192k'},
        'eac3': {'acodec': 'eac3', 'audio_bitrate': '192k'},
        'f32': {'acodec': 'pcm_f32le'},
        'f64': {'acodec': 'pcm_f64le'},
        's16': {'acodec': 'pcm_s16le'},
        's24': {'acodec': 'pcm_s24le'},
        's32': {'acodec': 'pcm_s32le'},
        'u8': {'acodec': 'pcm_u8'},
    }

    def convert(self) -> bytes:
        try:
            import ffmpeg
        except ImportError:
            raise RuntimeError("ffmpeg-python nÃ£o instalado")

        format_config = self.AUDIO_FORMATS.get(self.output_format.lower(), {})
        output_path = self._create_temp_file(f".{self.output_format}")

        stream = ffmpeg.input(self.input_path)
        stream = ffmpeg.output(stream, output_path, **format_config)
        ffmpeg.run(stream, overwrite_output=True, quiet=True)

        with open(output_path, 'rb') as f:
            return f.read()

class GenericVideoConverter(BaseConverter):
    """Conversor genÃ©rico de vÃ­deo para qualquer formato FFmpeg"""
    
    VIDEO_FORMATS = {
        'mp4': {'vcodec': 'libx264', 'acodec': 'aac', 'preset': 'medium'},
        'avi': {'vcodec': 'mpeg4', 'acodec': 'libmp3lame', 'q:v': '5'},
        'mkv': {'vcodec': 'libx264', 'acodec': 'aac', 'preset': 'medium'},
        'mov': {'vcodec': 'libx264', 'acodec': 'aac', 'preset': 'medium'},
        'flv': {'vcodec': 'mpeg4', 'acodec': 'libmp3lame'},
        'webm': {'vcodec': 'libvpx', 'acodec': 'libopus'},
        '3gp': {'vcodec': 'mpeg4', 'acodec': 'aac', 's': '320x240'},
        'wmv': {'vcodec': 'mpeg4', 'acodec': 'wmav2'},
        'asf': {'vcodec': 'mpeg4', 'acodec': 'wmav2'},
        'mod': {'vcodec': 'mpeg2video', 'acodec': 'mp2'},
        'mts': {'vcodec': 'mpeg2video', 'acodec': 'ac3'},
        'ts': {'vcodec': 'mpeg2video', 'acodec': 'ac3'},
        'vob': {'vcodec': 'mpeg2video', 'acodec': 'ac3', 'target': 'pal-dvd'},
        'm2ts': {'vcodec': 'mpeg2video', 'acodec': 'ac3'},
        'ogv': {'vcodec': 'libtheora', 'acodec': 'libvorbis'},
    }

    def convert(self) -> bytes:
        try:
            import ffmpeg
        except ImportError:
            raise RuntimeError("ffmpeg-python nÃ£o instalado")

        format_config = self.VIDEO_FORMATS.get(self.output_format.lower(), {})
        output_path = self._create_temp_file(f".{self.output_format}")

        stream = ffmpeg.input(self.input_path)
        stream = ffmpeg.output(stream, output_path, **format_config)
        ffmpeg.run(stream, overwrite_output=True, quiet=True)

        with open(output_path, 'rb') as f:
            return f.read()

class GenericImageConverter(BaseConverter):
    """Conversor genÃ©rico de imagem usando Pillow"""
    
    IMAGE_FORMATS = {
        'jpg': 'JPEG',
        'jpeg': 'JPEG',
        'png': 'PNG',
        'bmp': 'BMP',
        'gif': 'GIF',
        'tiff': 'TIFF',
        'ico': 'ICO',
        'webp': 'WEBP',
        'tga': 'TGA',
        'pnm': 'PPM',
        'ppm': 'PPM',
        'xbm': 'XBM',
        'xpm': 'XPM',
    }

    def convert(self) -> bytes:
        try:
            from PIL import Image
        except ImportError:
            raise RuntimeError("Pillow nÃ£o instalado")

        output_format = self.IMAGE_FORMATS.get(self.output_format.lower(), self.output_format.upper())
        output_path = self._create_temp_file(f".{self.output_format}")

        with Image.open(self.input_path) as img:
            # Converter para modo apropriado
            if output_format == 'JPEG':
                if img.mode in ('RGBA', 'LA', 'P'):
                    img = img.convert('RGB')
                img.save(output_path, output_format, quality=90)
            elif output_format in ('PNG', 'WEBP'):
                if img.mode == 'P':
                    img = img.convert('RGBA')
                img.save(output_path, output_format)
            else:
                img.save(output_path, output_format)

        with open(output_path, 'rb') as f:
            return f.read()

# ========== REGISTRO DE CONVERSORES ==========

from conversores import (
    BaseConverter,
    DocumentConverter,
    GenericAudioConverter,
    GenericImageConverter,
    GenericVideoConverter,
    SpreadsheetConverter,
    TextConverter,
    VideoToAudioConverter,
)

CONVERTERS = {
    # ========== ÃUDIO (20+ formatos) ==========
    ('audio', 'mp3'): GenericAudioConverter,
    ('audio', 'wav'): GenericAudioConverter,
    ('audio', 'flac'): GenericAudioConverter,
    ('audio', 'aac'): GenericAudioConverter,
    ('audio', 'm4a'): GenericAudioConverter,
    ('audio', 'ogg'): GenericAudioConverter,
    ('audio', 'opus'): GenericAudioConverter,
    ('audio', 'aiff'): GenericAudioConverter,
    ('audio', 'weba'): GenericAudioConverter,
    ('audio', 'ac3'): GenericAudioConverter,
    ('audio', 'dts'): GenericAudioConverter,
    ('audio', 'eac3'): GenericAudioConverter,
    ('audio', 'f32'): GenericAudioConverter,
    ('audio', 'f64'): GenericAudioConverter,
    ('audio', 's16'): GenericAudioConverter,
    ('audio', 's24'): GenericAudioConverter,
    ('audio', 's32'): GenericAudioConverter,
    ('audio', 'u8'): GenericAudioConverter,
    ('audio', 'wma'): GenericAudioConverter,
    ('audio', 'vorbis'): GenericAudioConverter,

    # ========== VÃDEO (20+ formatos) ==========
    ('video', 'mp4'): GenericVideoConverter,
    ('video', 'avi'): GenericVideoConverter,
    ('video', 'mkv'): GenericVideoConverter,
    ('video', 'mov'): GenericVideoConverter,
    ('video', 'flv'): GenericVideoConverter,
    ('video', 'webm'): GenericVideoConverter,
    ('video', '3gp'): GenericVideoConverter,
    ('video', 'wmv'): GenericVideoConverter,
    ('video', 'asf'): GenericVideoConverter,
    ('video', 'mod'): GenericVideoConverter,
    ('video', 'mts'): GenericVideoConverter,
    ('video', 'ts'): GenericVideoConverter,
    ('video', 'vob'): GenericVideoConverter,
    ('video', 'm2ts'): GenericVideoConverter,
    ('video', 'ogv'): GenericVideoConverter,
    ('video', 'wav'): VideoConverter,  # Extrair Ã¡udio
    ('video', 'm4v'): GenericVideoConverter,
    ('video', 'f4v'): GenericVideoConverter,
    ('video', 'insv'): GenericVideoConverter,
    ('video', 'qt'): GenericVideoConverter,
    ('video', 'wav'): VideoToAudioConverter,

    # ========== IMAGEM (20+ formatos) ==========
    ('image', 'jpg'): GenericImageConverter,
    ('image', 'jpeg'): GenericImageConverter,
    ('image', 'png'): GenericImageConverter,
    ('image', 'bmp'): GenericImageConverter,
    ('image', 'gif'): GenericImageConverter,
    ('image', 'tiff'): GenericImageConverter,
    ('image', 'ico'): GenericImageConverter,
    ('image', 'webp'): GenericImageConverter,
    ('image', 'tga'): GenericImageConverter,
    ('image', 'pnm'): GenericImageConverter,
    ('image', 'ppm'): GenericImageConverter,
    ('image', 'xbm'): GenericImageConverter,
    ('image', 'xpm'): GenericImageConverter,
    ('image', 'svg'): GenericImageConverter,
    ('image', 'cur'): GenericImageConverter,
    ('image', 'heic'): GenericImageConverter,
    ('image', 'heif'): GenericImageConverter,
    ('image', 'jfif'): GenericImageConverter,
    ('image', 'jp2'): GenericImageConverter,
    ('image', 'jp2k'): GenericImageConverter,

    # ========== DOCUMENTOS ==========
    ('document', 'pdf'): DocumentConverter,
    ('document', 'docx'): DocumentConverter,
    ('document', 'doc'): DocumentConverter,
    ('document', 'txt'): DocumentConverter,
    ('document', 'odt'): DocumentConverter,
    ('document', 'rtf'): DocumentConverter,

    # ========== PLANILHAS ==========
    ('spreadsheet', 'csv'): SpreadsheetConverter,
    ('spreadsheet', 'xlsx'): SpreadsheetConverter,
    ('spreadsheet', 'xls'): SpreadsheetConverter,
    ('spreadsheet', 'ods'): SpreadsheetConverter,
    ('spreadsheet', 'tsv'): SpreadsheetConverter,

    # ========== TEXTO ==========
    ('text', 'pdf'): TextConverter,
    ('text', 'txt'): TextConverter,
}

def get_converter(file_type: str, target_format: str, input_path: str) -> Optional[BaseConverter]:
    """Factory para obter conversor apropriado."""
    key = (file_type.lower(), target_format.lower())
    converter_class = CONVERTERS.get(key)
    if converter_class:
        return converter_class(input_path, target_format)
    return None

# ========== HELPERS REUTILIZADOS DO I'AM.pdf ==========

def file_sha1(path: str) -> str:
    """Calcula SHA1 do arquivo."""
    h = hashlib.sha1()
    with open(path, "rb") as f:
        for b in iter(lambda: f.read(1024 * 1024), b""):
            h.update(b)
    return h.hexdigest()

def cache_path(suffix: str, key: str) -> str:
    """Gera caminho de cache."""
    base = tempfile.gettempdir()
    fname = f"am_cache_{key}{suffix}"
    return os.path.join(base, fname)

def save_converted_copy(filename: str, data: bytes) -> None:
    """Salva cÃ³pia do arquivo convertido no diretÃ³rio exports."""
    try:
        os.makedirs(EXPORTS_DIR, exist_ok=True)
        safe = re.sub(r"[^A-Za-z0-9._-]+", "_", filename.strip()) or "convertido.bin"
        stem, ext = os.path.splitext(safe)
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
        target = os.path.join(EXPORTS_DIR, f"{stem}_{stamp}{ext}")
        with open(target, "wb") as f:
            f.write(data)
    except Exception:
        pass

def save_upload_to_temp(storage, per_file_limit_bytes: int = None) -> str:
    """Salva upload em arquivo temporÃ¡rio com limite de tamanho."""
    if per_file_limit_bytes is None:
        per_file_limit_bytes = config.PER_FILE_LIMIT

    fd, path = tempfile.mkstemp(); os.close(fd)
    wrote = 0
    try:
        storage.stream.seek(0)
    except Exception:
        pass
    with open(path, "wb") as f:
        while True:
            chunk = storage.stream.read(1024 * 1024)
            if not chunk:
                break
            wrote += len(chunk)
            if per_file_limit_bytes is not None and wrote > per_file_limit_bytes:
                try: os.remove(path)
                except: pass
                raise ValueError(f"Arquivo excede {config.PER_FILE_LIMIT_MB} MB.")
            f.write(chunk)
    return path

def save_upload_to_destination(storage, target_path: str, per_file_limit_bytes: int = None) -> str:
    """Salva upload em caminho fixo usando chunks para limitar uso de RAM."""
    if per_file_limit_bytes is None:
        per_file_limit_bytes = config.PER_FILE_LIMIT

    os.makedirs(os.path.dirname(target_path) or ".", exist_ok=True)
    wrote = 0
    try:
        storage.stream.seek(0)
    except Exception:
        pass

    with open(target_path, "wb") as handle:
        while True:
            chunk = storage.stream.read(1024 * 1024)
            if not chunk:
                break
            wrote += len(chunk)
            if per_file_limit_bytes is not None and wrote > per_file_limit_bytes:
                try:
                    os.remove(target_path)
                except OSError:
                    pass
                raise ValueError(f"Arquivo excede {config.PER_FILE_LIMIT_MB} MB.")
            handle.write(chunk)

    return target_path

def save_converted_file_copy(filename: str, source_path: str) -> Optional[str]:
    """Persiste arquivo convertido no diretÃ³rio exports usando cÃ³pia em chunks."""
    try:
        os.makedirs(EXPORTS_DIR, exist_ok=True)
        safe = re.sub(r"[^A-Za-z0-9._-]+", "_", filename.strip()) or "convertido.bin"
        stem, ext = os.path.splitext(safe)
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
        target = os.path.join(EXPORTS_DIR, f"{stem}_{stamp}{ext}")
        with open(source_path, "rb") as source, open(target, "wb") as destination:
            shutil.copyfileobj(source, destination, length=1024 * 1024)
        return target
    except Exception:
        logging.getLogger(__name__).exception("Falha ao salvar cÃ³pia convertida: %s", filename)
        return None

# ========== VALIDAÃ‡Ã•ES DE SEGURANÃ‡A ==========

def validate_file_type(filename: str) -> Tuple[str, str]:
    """Valida tipo de arquivo baseado na extensÃ£o."""
    ext = filename.lower().split('.')[-1]

    # Mapeamento de extensÃµes para tipos
    type_map = {
        # Ãudio
        'ogg': 'audio',
        'mp3': 'audio',
        'wav': 'audio',

        # VÃ­deo
        'mp4': 'video',
        'avi': 'video',
        'm4v': 'video',

        # Imagem
        'png': 'image',
        'jpg': 'image',
        'jpeg': 'image',
        'gif': 'image',
        'bmp': 'image',
        'tiff': 'image',
        'webp': 'image',

        # Documentos
        'docx': 'document',
        'doc': 'document',
        'pdf': 'document',

        # Texto
        'txt': 'text',

        # Planilhas
        'xlsx': 'spreadsheet',
        'xls': 'spreadsheet',
        'csv': 'spreadsheet',
    }

    file_type = type_map.get(ext)
    if not file_type:
        raise ValueError(f"Tipo de arquivo nÃ£o suportado: .{ext}")

    return file_type, ext

def sanitize_filename(filename: str) -> str:
    """Sanitiza nome do arquivo."""
    return re.sub(r"[^A-Za-z0-9._-]+", "_", filename.strip()) or "arquivo"

# ========== VARIÃVEIS GLOBAIS ==========
EXPORTS_DIR = config.EXPORTS_DIR
PER_FILE_LIMIT_MB = config.PER_FILE_LIMIT_MB

# ========== IMPORTAR ROTAS ==========
# Importar rotas para registrar no app (FORA do if __name__ para funcionar com WSGI)
try:
    from rotas.app_routes import *
except ImportError as e:
    print(f"âš ï¸  Aviso ao importar rotas: {e}")

# ========== INICIALIZAÃ‡ÃƒO DA FILA v2.2.0 ==========
from queue_manager import QueueManager
from worker import get_worker, init_worker, start_worker, stop_worker
from websocket_manager import get_ws_manager
from rotas.queue_routes import register_queue_routes

# Inicializar gerenciadores de fila
queue_manager = QueueManager()
ws_manager = get_ws_manager()

def create_conversion_func():
    """
    Factory para criar funÃ§Ã£o de conversÃ£o que atualiza progresso.
    Esta funÃ§Ã£o serÃ¡ chamada pelo worker para processar tarefas da fila.
    """
    def convert_with_progress(task_id: str, task: dict):
        """
        Converte arquivo e atualiza progresso via WebSocket.

        Args:
            task_id: ID da tarefa
            task: Dict com dados da tarefa (filename_original, target_format, etc.)
        """
        try:
            filename = task['filename_original']
            target_format = task['target_format']
            file_path = os.path.join(config.UPLOAD_FOLDER, filename)
            logger = logging.getLogger(__name__)

            # Verificar se arquivo existe
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"Arquivo nÃ£o encontrado: {file_path}")

            # Enviar progresso 10% - Iniciando
            ws_manager.send_progress(task_id, 10, 30)

            # Validar tipo do arquivo
            file_type, ext = validate_file_type(filename)
            logger.debug("Tipo de arquivo validado | task_id=%s tipo=%s ext=%s", task_id, file_type, ext)

            # Enviar progresso 20% - ValidaÃ§Ã£o OK
            ws_manager.send_progress(task_id, 20, 25)

            # Obter conversor apropriado
            converter = get_converter(file_type, target_format, file_path)
            logger.debug("Conversor obtido | task_id=%s converter=%s", task_id, type(converter).__name__ if converter else None)
            if not converter:
                raise ValueError(f"ConversÃ£o {file_type} â†’ {target_format} nÃ£o suportada")

            # Enviar progresso 30% - Conversor preparado
            ws_manager.send_progress(task_id, 30, 20)

            # Realizar conversÃ£o
            logger.debug("Iniciando conversão | task_id=%s", task_id)
            with converter:
                converted_path = converter.convert()
                converter.temp_files = []
            export_filename = converter.get_output_filename(sanitize_filename(filename))
            export_path = save_converted_file_copy(export_filename, converted_path)
            if not export_path:
                raise RuntimeError("Falha ao persistir arquivo convertido")
            try:
                os.remove(converted_path)
            except OSError:
                pass

            ws_manager.send_progress(task_id, 80, 10)
            ws_manager.send_progress(task_id, 100, 0)
            result = {"filename": os.path.basename(export_path)}
            ws_manager.send_completion(task_id, True, result=result)
            return True, result, None

        except Exception as e:
            logger = logging.getLogger(__name__)
            ws_manager.send_completion(task_id, False, error=str(e))
            logger.error(f"Erro na conversÃ£o {task_id}: {e}", exc_info=True)
            return False, None, str(e)

    return convert_with_progress

# Inicializar worker com funÃ§Ã£o de conversÃ£o
worker = init_worker(
    conversion_func=create_conversion_func(),
    queue_manager=queue_manager,
    max_concurrent=config.QUEUE_MAX_CONCURRENT,
    check_interval=config.QUEUE_CHECK_INTERVAL
)

register_queue_routes(app, conversion_func=create_conversion_func(), upload_folder=config.UPLOAD_FOLDER)

if sock is not None:
    @sock.route("/ws/queue/<task_id>")
    def queue_status_socket(ws, task_id):
        """Canal WebSocket para progresso em tempo real da fila."""
        ws_manager.register_connection(task_id, ws)
        try:
            task = queue_manager.get_status(task_id)
            if task:
                ws.send(json.dumps({
                    "type": "status",
                    "task_id": task_id,
                    "status": task["status"],
                    "progress": task["progress"],
                    "eta_seconds": task["eta_seconds"],
                    "result_filename": task["filename_result"],
                    "error": task["error_message"],
                }))

            while True:
                message = ws.receive()
                if message is None:
                    break
        except Exception as exc:
            logging.getLogger(__name__).warning("WebSocket encerrado | task_id=%s erro=%s", task_id, exc)
        finally:
            ws_manager.unregister_connection(task_id, ws)


@app.before_request
def ensure_background_services():
    """Garante que o worker esteja rodando antes de atender requests."""
    active_worker = get_worker()
    if active_worker and not active_worker.is_running:
        start_worker()

if __name__ == "__main__":
    try:
        app.run(debug=getattr(config, "DEBUG", False), host=config.HOST, port=config.PORT)
    finally:
        stop_worker()
        print("Worker parado")



