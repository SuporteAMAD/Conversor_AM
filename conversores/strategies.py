import csv
import logging
import os
import shutil
from pathlib import Path

from .base import BaseConverter

logger = logging.getLogger(__name__)


class GenericAudioConverter(BaseConverter):
    """Estratégia genérica de áudio baseada em FFmpeg."""

    AUDIO_FORMATS = {
        "mp3": {"acodec": "mp3", "audio_bitrate": "192k"},
        "wav": {"acodec": "pcm_s16le", "ar": "44100"},
        "flac": {"acodec": "flac"},
        "aac": {"acodec": "aac", "audio_bitrate": "192k"},
        "m4a": {"acodec": "aac", "audio_bitrate": "192k"},
        "ogg": {"acodec": "libvorbis", "audio_bitrate": "192k"},
        "opus": {"acodec": "libopus", "audio_bitrate": "128k"},
        "aiff": {"acodec": "pcm_s16le", "f": "aiff"},
        "weba": {"acodec": "libopus", "audio_bitrate": "128k"},
        "ac3": {"acodec": "ac3", "audio_bitrate": "192k"},
        "dts": {"acodec": "dts", "audio_bitrate": "192k"},
        "eac3": {"acodec": "eac3", "audio_bitrate": "192k"},
        "f32": {"acodec": "pcm_f32le"},
        "f64": {"acodec": "pcm_f64le"},
        "s16": {"acodec": "pcm_s16le"},
        "s24": {"acodec": "pcm_s24le"},
        "s32": {"acodec": "pcm_s32le"},
        "u8": {"acodec": "pcm_u8"},
        "wma": {"acodec": "wmav2"},
        "vorbis": {"acodec": "libvorbis", "audio_bitrate": "192k"},
    }

    def convert(self) -> str:
        try:
            import ffmpeg
        except ImportError as exc:
            raise RuntimeError("ffmpeg-python não instalado. Instale com: pip install ffmpeg-python") from exc

        format_config = self.AUDIO_FORMATS.get(self.output_format, {})
        output_path = self._create_temp_file(f".{self.output_format}")

        try:
            stream = ffmpeg.input(self.input_path)
            stream = ffmpeg.output(stream, output_path, **format_config)
            ffmpeg.run(stream, overwrite_output=True, quiet=True)
            return output_path
        except Exception:
            logger.exception("Falha na conversão de áudio para %s", self.output_format)
            raise


class VideoToAudioConverter(BaseConverter):
    """Estratégia específica para extrair áudio de vídeo."""

    def convert(self) -> str:
        try:
            import ffmpeg
        except ImportError as exc:
            raise RuntimeError("ffmpeg-python não instalado. Instale com: pip install ffmpeg-python") from exc

        output_path = self._create_temp_file(".wav")

        try:
            stream = ffmpeg.input(self.input_path)
            stream = ffmpeg.output(stream, output_path, acodec="pcm_s16le", ar="44100")
            ffmpeg.run(stream, overwrite_output=True, quiet=True)
            return output_path
        except Exception:
            logger.exception("Falha ao extrair áudio do vídeo %s", self.input_path)
            raise


class GenericVideoConverter(BaseConverter):
    """Estratégia genérica de vídeo baseada em FFmpeg."""

    VIDEO_FORMATS = {
        "mp4": {"vcodec": "libx264", "acodec": "aac", "preset": "medium"},
        "avi": {"vcodec": "mpeg4", "acodec": "libmp3lame", "q:v": "5"},
        "mkv": {"vcodec": "libx264", "acodec": "aac", "preset": "medium"},
        "mov": {"vcodec": "libx264", "acodec": "aac", "preset": "medium"},
        "flv": {"vcodec": "mpeg4", "acodec": "libmp3lame"},
        "webm": {"vcodec": "libvpx", "acodec": "libopus"},
        "3gp": {"vcodec": "mpeg4", "acodec": "aac", "s": "320x240"},
        "wmv": {"vcodec": "mpeg4", "acodec": "wmav2"},
        "asf": {"vcodec": "mpeg4", "acodec": "wmav2"},
        "mod": {"vcodec": "mpeg2video", "acodec": "mp2"},
        "mts": {"vcodec": "mpeg2video", "acodec": "ac3"},
        "ts": {"vcodec": "mpeg2video", "acodec": "ac3"},
        "vob": {"vcodec": "mpeg2video", "acodec": "ac3", "target": "pal-dvd"},
        "m2ts": {"vcodec": "mpeg2video", "acodec": "ac3"},
        "ogv": {"vcodec": "libtheora", "acodec": "libvorbis"},
        "m4v": {"vcodec": "libx264", "acodec": "aac", "preset": "medium"},
        "f4v": {"vcodec": "libx264", "acodec": "aac", "preset": "medium"},
        "insv": {"vcodec": "libx264", "acodec": "aac", "preset": "medium"},
        "qt": {"vcodec": "libx264", "acodec": "aac", "preset": "medium"},
    }

    def convert(self) -> str:
        try:
            import ffmpeg
        except ImportError as exc:
            raise RuntimeError("ffmpeg-python não instalado. Instale com: pip install ffmpeg-python") from exc

        format_config = self.VIDEO_FORMATS.get(self.output_format, {})
        output_path = self._create_temp_file(f".{self.output_format}")

        try:
            stream = ffmpeg.input(self.input_path)
            stream = ffmpeg.output(stream, output_path, **format_config)
            ffmpeg.run(stream, overwrite_output=True, quiet=True)
            return output_path
        except Exception:
            logger.exception("Falha na conversão de vídeo para %s", self.output_format)
            raise


class GenericImageConverter(BaseConverter):
    """Estratégia genérica de imagem usando Pillow."""

    IMAGE_FORMATS = {
        "jpg": "JPEG",
        "jpeg": "JPEG",
        "png": "PNG",
        "bmp": "BMP",
        "gif": "GIF",
        "tiff": "TIFF",
        "ico": "ICO",
        "webp": "WEBP",
        "tga": "TGA",
        "pnm": "PPM",
        "ppm": "PPM",
        "xbm": "XBM",
        "xpm": "XPM",
    }

    def convert(self) -> str:
        try:
            from PIL import Image
        except ImportError as exc:
            raise RuntimeError("Pillow não instalado. Instale com: pip install Pillow") from exc

        output_format = self.IMAGE_FORMATS.get(self.output_format, self.output_format.upper())
        output_path = self._create_temp_file(f".{self.output_format}")

        try:
            with Image.open(self.input_path) as img:
                if output_format == "JPEG":
                    if img.mode in ("RGBA", "LA", "P"):
                        img = img.convert("RGB")
                    img.save(output_path, output_format, quality=90)
                elif output_format in ("PNG", "WEBP"):
                    if img.mode == "P":
                        img = img.convert("RGBA")
                    img.save(output_path, output_format)
                else:
                    img.save(output_path, output_format)
            return output_path
        except Exception:
            logger.exception("Falha na conversão de imagem para %s", self.output_format)
            raise


class DocumentConverter(BaseConverter):
    """Estratégia para conversões documentais suportadas localmente."""

    def convert(self) -> str:
        try:
            from docx import Document
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
        except ImportError as exc:
            raise RuntimeError("Dependências documentais não instaladas") from exc

        input_ext = Path(self.input_path).suffix.lower()

        try:
            if self.output_format == "pdf" and input_ext == ".docx":
                doc = Document(self.input_path)
                output_path = self._create_temp_file(".pdf")
                pdf = SimpleDocTemplate(output_path, pagesize=letter)
                styles = getSampleStyleSheet()
                story = []

                for para in doc.paragraphs:
                    if para.text.strip():
                        story.append(Paragraph(para.text, styles["Normal"]))
                        story.append(Spacer(1, 12))

                pdf.build(story)
                return output_path

            if self.output_format == "docx" and input_ext == ".pdf":
                try:
                    from pdf2docx import Converter
                except ImportError as exc:
                    raise RuntimeError("pdf2docx não instalado. Instale com: pip install pdf2docx") from exc

                output_path = self._create_temp_file(".docx")
                converter = Converter(self.input_path)
                try:
                    converter.convert(output_path, start=0, end=None)
                finally:
                    converter.close()
                return output_path

            if self.output_format == "txt" and input_ext == ".docx":
                doc = Document(self.input_path)
                output_path = self._create_temp_file(".txt")
                with open(output_path, "w", encoding="utf-8", newline="") as handle:
                    for para in doc.paragraphs:
                        if para.text.strip():
                            handle.write(para.text)
                            handle.write("\n")
                return output_path
        except Exception:
            logger.exception(
                "Falha na conversão documental de %s para %s",
                input_ext,
                self.output_format,
            )
            raise

        raise ValueError(f"Conversão documental não suportada: {input_ext} -> {self.output_format}")


class TextConverter(BaseConverter):
    """Estratégia para texto, priorizando leitura iterativa."""

    ENCODINGS_TO_TRY = ("utf-8", "utf-16", "utf-16-le", "utf-16-be", "latin-1", "cp1252")

    def _iter_text_lines(self):
        for encoding in self.ENCODINGS_TO_TRY:
            try:
                with open(self.input_path, "r", encoding=encoding) as handle:
                    for line in handle:
                        yield line
                return
            except UnicodeDecodeError:
                continue

        raise ValueError("Não foi possível decodificar o arquivo de texto")

    def convert(self) -> str:
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
        except ImportError as exc:
            raise RuntimeError("reportlab não instalado. Instale com: pip install reportlab") from exc

        try:
            if self.output_format == "txt":
                output_path = self._create_temp_file(".txt")
                shutil.copyfile(self.input_path, output_path)
                return output_path

            if self.output_format == "pdf":
                output_path = self._create_temp_file(".pdf")
                pdf = SimpleDocTemplate(output_path, pagesize=letter)
                styles = getSampleStyleSheet()
                story = []

                for line in self._iter_text_lines():
                    clean_line = line.strip()
                    if clean_line:
                        story.append(Paragraph(clean_line, styles["Normal"]))
                        story.append(Spacer(1, 6))

                pdf.build(story)
                return output_path
        except Exception:
            logger.exception("Falha na conversão de texto para %s", self.output_format)
            raise

        raise ValueError(f"Conversão de texto não suportada para {self.output_format}")


class SpreadsheetConverter(BaseConverter):
    """Estratégia para planilhas com foco em CSV/TSV/XLSX."""

    def convert(self) -> str:
        input_ext = Path(self.input_path).suffix.lower()

        try:
            if self.output_format in {"csv", "tsv"}:
                delimiter = "," if self.output_format == "csv" else "\t"
                output_path = self._create_temp_file(f".{self.output_format}")

                if input_ext in {".csv", ".tsv"}:
                    input_delimiter = "," if input_ext == ".csv" else "\t"
                    with open(self.input_path, "r", encoding="utf-8", newline="") as source:
                        reader = csv.reader(source, delimiter=input_delimiter)
                        with open(output_path, "w", encoding="utf-8", newline="") as target:
                            writer = csv.writer(target, delimiter=delimiter)
                            for row in reader:
                                writer.writerow(row)
                    return output_path

                if input_ext in {".xlsx", ".xlsm"}:
                    from openpyxl import load_workbook

                    workbook = load_workbook(self.input_path, read_only=True, data_only=True)
                    try:
                        worksheet = workbook.active
                        with open(output_path, "w", encoding="utf-8", newline="") as target:
                            writer = csv.writer(target, delimiter=delimiter)
                            for row in worksheet.iter_rows(values_only=True):
                                writer.writerow(["" if value is None else value for value in row])
                    finally:
                        workbook.close()
                    return output_path

            if self.output_format == "xlsx":
                from openpyxl import Workbook, load_workbook

                output_path = self._create_temp_file(".xlsx")
                workbook = Workbook(write_only=True)
                worksheet = workbook.create_sheet()

                try:
                    if input_ext in {".csv", ".tsv"}:
                        input_delimiter = "," if input_ext == ".csv" else "\t"
                        with open(self.input_path, "r", encoding="utf-8", newline="") as source:
                            reader = csv.reader(source, delimiter=input_delimiter)
                            for row in reader:
                                worksheet.append(row)
                    elif input_ext in {".xlsx", ".xlsm"}:
                        source_wb = load_workbook(self.input_path, read_only=True, data_only=True)
                        try:
                            source_ws = source_wb.active
                            for row in source_ws.iter_rows(values_only=True):
                                worksheet.append(list(row))
                        finally:
                            source_wb.close()
                    else:
                        raise ValueError(f"Formato de entrada de planilha não suportado: {input_ext}")

                    workbook.save(output_path)
                    return output_path
                finally:
                    workbook.close()
        except Exception:
            logger.exception(
                "Falha na conversão de planilha de %s para %s",
                input_ext,
                self.output_format,
            )
            raise

        raise ValueError(f"Conversão de planilha não suportada: {input_ext} -> {self.output_format}")
